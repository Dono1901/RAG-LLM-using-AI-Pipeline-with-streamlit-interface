"""
Ingestion Pipeline for Financial Documents.

Orchestrates: parse → chunk → embed → index.
Supports Excel (.xlsx, .xlsm, .xls, .csv) and PDF files.

Produces RAGChunk objects with embeddings ready for vector storage.
Integrates with SimpleRAG's document/embedding stores.
"""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

from config import settings
from document_chunker import RAGChunk, chunk_excel_sheet, chunk_table, chunk_text_content

logger = logging.getLogger(__name__)

# Supported file extensions
EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xls", ".csv", ".tsv"}
PDF_EXTENSIONS = {".pdf"}
TEXT_EXTENSIONS = {".txt", ".md", ".docx"}


def _detect_sheet_section_type(df: pd.DataFrame, sheet_name: str) -> str:
    """Detect the financial section type of an Excel sheet.

    Uses column names, row labels, and sheet name to classify.
    """
    from line_item_mapper import map_label

    # Check sheet name first
    sheet_lower = sheet_name.lower()
    # Ordered list of (hint, section_type) — longer/more specific hints first
    sheet_type_hints = [
        ("income", "income_statement"),
        ("p&l", "income_statement"),
        ("profit", "income_statement"),
        ("revenue", "income_statement"),
        ("balance", "balance_sheet"),
        ("assets", "balance_sheet"),
        ("cash flow", "cash_flow_statement"),
        ("cash_flow", "cash_flow_statement"),
        ("equity", "equity_statement"),
        ("budget", "budget"),
        ("forecast", "forecast"),
        ("fcst", "forecast"),
        ("kpi", "kpi_dashboard"),
        ("metric", "kpi_dashboard"),
        ("dashboard", "kpi_dashboard"),
        ("debt", "debt_schedule"),
        ("loan", "debt_schedule"),
        ("covenant", "covenant_analysis"),
        ("280e", "280e_tax"),
        ("production", "production_schedule"),
        ("cultivation", "production_schedule"),
        ("rent roll", "rent_roll"),
        ("construction", "construction_budget"),
        ("cap rate", "cap_rate_analysis"),
        ("dcf", "dcf"),
        ("lbo", "lbo"),
        ("valuation", "dcf"),
        ("comps", "comparable_companies"),
        ("waterfall", "debt_schedule"),
        ("capital account", "capital_account"),
        ("fund", "fund_performance"),
        ("assumptions", "assumptions"),
        ("sensitivity", "sensitivity_analysis"),
        ("scenario", "scenario_analysis"),
    ]

    for hint, section in sheet_type_hints:
        if hint in sheet_lower:
            return section

    # Sample row labels to detect content type
    label_col = _find_label_column(df)
    if label_col is not None:
        labels = df.iloc[:30, label_col].dropna().astype(str).tolist()
        category_counts: Dict[str, int] = {}
        for label in labels:
            mapped = map_label(label)
            if mapped:
                cat = mapped.category
                category_counts[cat] = category_counts.get(cat, 0) + 1

        if category_counts:
            top_category = max(category_counts, key=lambda k: category_counts[k])
            # Map category to section type
            cat_to_section = {
                "income_statement": "income_statement",
                "balance_sheet": "balance_sheet",
                "cash_flow": "cash_flow_statement",
                "valuation": "dcf",
                "debt": "debt_schedule",
                "cannabis": "280e_tax",
                "cash_forecast": "forecast",
                "kpi": "kpi_dashboard",
                "saas": "kpi_dashboard",
                "construction": "construction_budget",
                "fund": "fund_performance",
            }
            return cat_to_section.get(top_category, "general")

    return "general"


def _find_label_column(df: pd.DataFrame) -> Optional[int]:
    """Find the column most likely to contain line item labels.

    Often column B (index 1) in real models, sometimes column A.
    Looks for the column with the most string values.
    """
    if df.empty:
        return None

    best_col = None
    best_score = 0

    for col_idx in range(min(5, len(df.columns))):  # Check first 5 columns
        col = df.iloc[:, col_idx]
        str_count = col.apply(lambda x: isinstance(x, str) and len(str(x).strip()) > 2).sum()
        # Penalize columns that are mostly numeric
        num_count = pd.to_numeric(col, errors="coerce").notna().sum()
        score = str_count - num_count * 0.5
        if score > best_score:
            best_score = score
            best_col = col_idx

    return best_col if best_score > 3 else None


def _df_to_markdown(df: pd.DataFrame, max_rows: int = 200) -> str:
    """Convert a DataFrame to markdown table format.

    Limits to max_rows to prevent extremely large chunks.
    """
    if df.empty:
        return ""

    # Truncate if needed
    truncated = df.head(max_rows)

    try:
        return truncated.to_markdown(index=False)
    except Exception:
        # Fallback: simple pipe-delimited format
        lines = []
        headers = [str(c) for c in truncated.columns]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for _, row in truncated.iterrows():
            vals = [str(v) if pd.notna(v) else "" for v in row]
            lines.append("| " + " | ".join(vals) + " |")
        return "\n".join(lines)


def ingest_excel(
    file_path: Path,
    child_token_target: int = 250,
    parent_token_target: int = 1200,
) -> List[RAGChunk]:
    """Ingest an Excel file into RAGChunks.

    Processes each sheet separately, detects section types,
    and creates financial-aware chunks.

    Args:
        file_path: Path to the Excel file.
        child_token_target: Target child chunk size in tokens.
        parent_token_target: Target parent chunk size in tokens.

    Returns:
        List of RAGChunk objects ready for embedding.
    """
    file_path = Path(file_path)
    source = file_path.name
    all_chunks: List[RAGChunk] = []

    try:
        suffix = file_path.suffix.lower()
        if suffix in (".csv", ".tsv"):
            sep = "\t" if suffix == ".tsv" else ","
            df = pd.read_csv(file_path, sep=sep, nrows=settings.max_workbook_rows)
            # Fix unnamed columns from headerless Excel exports
            df.columns = [
                c if not str(c).startswith("Unnamed") else f"Col_{i}"
                for i, c in enumerate(df.columns)
            ]
            sheets = [("Sheet1", df)]
        else:
            xls = pd.ExcelFile(file_path)
            sheets = []
            for sheet_name in xls.sheet_names:
                try:
                    df = pd.read_excel(xls, sheet_name=sheet_name, nrows=settings.max_workbook_rows)
                    # Fix unnamed columns
                    df.columns = [
                        c if not str(c).startswith("Unnamed") else f"Col_{i}"
                        for i, c in enumerate(df.columns)
                    ]
                    # Skip empty sheets
                    if df.empty or (df.shape[0] < 2 and df.shape[1] < 2):
                        continue
                    sheets.append((sheet_name, df))
                except Exception as e:
                    logger.warning("Failed to read sheet '%s' from %s: %s", sheet_name, source, e)

        for sheet_name, df in sheets:
            section_type = _detect_sheet_section_type(df, sheet_name)
            md = _df_to_markdown(df)

            if not md.strip():
                continue

            meta = {
                "source_file": source,
                "sheet_name": sheet_name,
                "file_type": "excel",
                "section_type": section_type,
                "row_count": len(df),
                "col_count": len(df.columns),
            }

            sheet_chunks = chunk_excel_sheet(
                md,
                source=source,
                sheet_name=sheet_name,
                section_type=section_type,
                metadata=meta,
            )
            all_chunks.extend(sheet_chunks)

        logger.info(
            "Ingested Excel '%s': %d sheets -> %d chunks",
            source, len(sheets), len(all_chunks),
        )

    except Exception as e:
        logger.error("Failed to ingest Excel file %s: %s", source, e)

    return all_chunks


def ingest_pdf(
    file_path: Path,
    child_token_target: int = 250,
    parent_token_target: int = 1200,
) -> List[RAGChunk]:
    """Ingest a PDF file into RAGChunks.

    Uses pdf_parser for structured extraction, then document_chunker
    for financial-aware chunking.

    Args:
        file_path: Path to the PDF file.
        child_token_target: Target child chunk size in tokens.
        parent_token_target: Target parent chunk size in tokens.

    Returns:
        List of RAGChunk objects ready for embedding.
    """
    from pdf_parser import parse_pdf

    file_path = Path(file_path)
    source = file_path.name
    all_chunks: List[RAGChunk] = []

    try:
        parsed = parse_pdf(file_path)

        doc_meta = {
            "source_file": source,
            "file_type": "pdf",
            "company": parsed.company,
            "period": parsed.period,
            "total_pages": parsed.total_pages,
            **parsed.metadata,
        }

        for section in parsed.sections:
            # Handle tables as atomic chunks
            for table in section.tables:
                table_chunk = chunk_table(
                    table,
                    source=source,
                    section_type=section.section_type,
                    section_title=section.title,
                    page_start=section.page_start,
                    page_end=section.page_end,
                    metadata=doc_meta,
                )
                all_chunks.append(table_chunk)

            # Chunk the non-table text content
            # Remove tables from section content to avoid double-indexing
            text_content = section.content
            for table in section.tables:
                text_content = text_content.replace(table, "")

            text_content = text_content.strip()
            if text_content and len(text_content) > 50:
                text_chunks = chunk_text_content(
                    text_content,
                    source=source,
                    section_type=section.section_type,
                    section_title=section.title,
                    page_start=section.page_start,
                    page_end=section.page_end,
                    child_token_target=child_token_target,
                    parent_token_target=parent_token_target,
                    metadata=doc_meta,
                )
                all_chunks.extend(text_chunks)

        logger.info(
            "Ingested PDF '%s': %d sections -> %d chunks",
            source, len(parsed.sections), len(all_chunks),
        )

    except Exception as e:
        logger.error("Failed to ingest PDF file %s: %s", source, e)

    return all_chunks


def ingest_text(
    file_path: Path,
    child_token_target: int = 250,
    parent_token_target: int = 1200,
) -> List[RAGChunk]:
    """Ingest a text/markdown/docx file into RAGChunks.

    Args:
        file_path: Path to the text file.
        child_token_target: Target child chunk size.
        parent_token_target: Target parent chunk size.

    Returns:
        List of RAGChunk objects.
    """
    file_path = Path(file_path)
    source = file_path.name
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = file_path.read_text(encoding="utf-8")

        if not text.strip():
            return []

        meta = {
            "source_file": source,
            "file_type": "text",
        }

        chunks = chunk_text_content(
            text,
            source=source,
            section_type="general",
            section_title=file_path.stem,
            child_token_target=child_token_target,
            parent_token_target=parent_token_target,
            metadata=meta,
        )

        logger.info("Ingested text '%s': %d chunks", source, len(chunks))
        return chunks

    except Exception as e:
        logger.error("Failed to ingest text file %s: %s", source, e)
        return []


def ingest_file(file_path: Path) -> List[RAGChunk]:
    """Ingest a single file into RAGChunks.

    Auto-detects file type and dispatches to the appropriate handler.

    Args:
        file_path: Path to the file.

    Returns:
        List of RAGChunk objects ready for embedding.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix in EXCEL_EXTENSIONS:
        return ingest_excel(file_path)
    elif suffix in PDF_EXTENSIONS:
        return ingest_pdf(file_path)
    elif suffix in TEXT_EXTENSIONS:
        return ingest_text(file_path)
    else:
        logger.warning("Unsupported file type: %s", suffix)
        return []


def chunks_to_documents(chunks: List[RAGChunk]) -> List[Dict[str, Any]]:
    """Convert RAGChunks to the document dict format used by SimpleRAG.

    The text used for embedding includes the NL description prefix
    when available, improving retrieval quality for numeric content.

    Args:
        chunks: List of RAGChunk objects.

    Returns:
        List of document dicts compatible with SimpleRAG.documents.
    """
    documents: List[Dict[str, Any]] = []

    for chunk in chunks:
        # Build the text to embed: NL description + actual content
        embed_text = chunk.text
        if chunk.nl_description:
            embed_text = chunk.nl_description + "\n\n" + chunk.text

        doc = {
            "source": chunk.source,
            "content": embed_text,
            "type": chunk.metadata.get("file_type", "unknown"),
            "metadata": {
                "chunk_id": chunk.chunk_id,
                "parent_id": chunk.parent_id,
                "parent_text": chunk.parent_text,
                "section_type": chunk.section_type,
                "section_title": chunk.section_title,
                "page_start": chunk.page_start,
                "page_end": chunk.page_end,
                "is_table": chunk.is_table,
                "source": chunk.source,
                **chunk.metadata,
            },
        }
        documents.append(doc)

    return documents
